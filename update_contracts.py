from pathlib import Path
from docx import Document

INDIVIDUAL_NAME = "岸谷 祥平"
INDIVIDUAL_ADDRESS = "〒141-0001 東京都品川区北品川5-3-1-3609"


def replace_placeholder(doc: Document, placeholder: str, replacement: str):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement)


def update_kou_section(doc: Document):
    """Find paragraph starting with '甲：' and set subsequent lines for address and name."""
    paragraphs = doc.paragraphs
    for idx, p in enumerate(paragraphs):
        if p.text.strip().startswith('甲：'):
            # Update the paragraph with address
            p.text = f'甲：{INDIVIDUAL_ADDRESS}'
            # Ensure next paragraph is the individual's name
            if idx + 1 < len(paragraphs):
                paragraphs[idx + 1].text = INDIVIDUAL_NAME
            # Optional: clear further blank lines if any
            if idx + 2 < len(paragraphs):
                paragraphs[idx + 2].text = ''
            break


def process_file(path: Path):
    doc = Document(path)
    # Replace placeholder '●●' with individual's name
    replace_placeholder(doc, '●●', INDIVIDUAL_NAME)
    # Update 甲 section
    update_kou_section(doc)
    doc.save(path)
    print(f'Updated {path}')


if __name__ == '__main__':
    for file_path in Path('docs').glob('*.docx'):
        process_file(file_path) 